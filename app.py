from flask import Flask, render_template, request, send_file, jsonify
import spacy
import pdfplumber
import docx 
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import re
import csv
import os
import google.generativeai as genai
from kafka_utils.kafka_producer import send_to_kafka
app = Flask(__name__)
nlp = spacy.load("en_core_web_sm")

# Gemini API Configuration
genai.configure(api_key="")
model = genai.GenerativeModel("gemini-1.5-flash")

global_results = []

def extract_text_from_pdf(pdf_path):
    with pdfplumber.open(pdf_path) as pdf:
        text = ""
        for page in pdf.pages:
            text += page.extract_text()
        return text

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\n"
    return text

def extract_entities(text):
    emails = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
    print("Emails found: ", emails)
    name_matches = re.findall(r'\b([A-Z][a-z]+(?: [A-Z][a-z]+)+)\b', text)
    names = name_matches if name_matches else ["Name"]
    emails = emails if emails else ["Email"]
    return emails, names

@app.route('/', methods=['GET', 'POST'])
def index():
    global global_results
    global_results = []

    if request.method == 'POST':
        job_description = request.form['job_description']
        resume_files = request.files.getlist('resume_files')

        if not os.path.exists("uploads"):
            os.makedirs("uploads")

        processed_resumes = []
        for resume_file in resume_files:
            resume_path = os.path.join("uploads", resume_file.filename)
            resume_file.save(resume_path)

            if resume_path.lower().endswith(".pdf"):
                resume_text = extract_text_from_pdf(resume_path)
            elif resume_path.lower().endswith(".docx"):
                resume_text = extract_text_from_docx(resume_path)
            else:
                resume_text = ""

            emails, names = extract_entities(resume_text)
            processed_resumes.append((names, emails, resume_text))

        tfidf_vectorizer = TfidfVectorizer()
        job_desc_vector = tfidf_vectorizer.fit_transform([job_description])

        shortlisted_cand = []
        for (names, emails, resume_text) in processed_resumes:
            resume_vector = tfidf_vectorizer.transform([resume_text])
            similarity = cosine_similarity(job_desc_vector, resume_vector)[0][0] * 100
            shortlisted_cand.append((names, emails, similarity))

            # Kafka call to send email notification
            if emails and names:
                email = emails[0]
                name = names[0]
                send_to_kafka(email, name, similarity, "Company_XYZ")

        shortlisted_cand.sort(key=lambda x: x[2], reverse=True)
        global_results = shortlisted_cand

    return render_template('index.html', results=global_results)

@app.route('/download_csv')
def download_csv():
    if not global_results:
        return "No data to download", 400

    csv_filename = "shortlisted_cand.csv"
    csv_full_path = os.path.join(os.path.abspath(os.path.dirname(__file__)), csv_filename)

    with open(csv_full_path, "w", newline="") as csv_file:
        writer = csv.writer(csv_file)
        writer.writerow(["Rank", "Name", "Email", "Similarity"])

        for rank, (names, emails, similarity) in enumerate(global_results, start=1):
            name = names[0] if names else "Name not extracted"
            email = emails[0] if emails else "Email not extracted"
            writer.writerow([rank, name, email, similarity])

    return send_file(csv_full_path, as_attachment=True, download_name="shortlisted_cand.csv")

@app.route('/gemini-chat', methods=['POST'])
def gemini_chat():
    user_input = request.json.get('message')
    if not user_input:
        return jsonify({'response': 'No input received'}), 400

    try:
        response = model.generate_content(user_input)
        return jsonify({'response': response.text})
    except Exception as e:
        return jsonify({'response': f"Error: {str(e)}"}), 500

if __name__ == '__main__':
    app.run(debug=True)
